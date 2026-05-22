import os
import zipfile
import shutil
import tempfile
import xml.etree.ElementTree as ET
from datetime import datetime

from docx import Document


class DOCXHandler:
    FILESYSTEM_EDITABLE_FIELDS = [
        "FileModifyDate",
        "FileAccessDate",
    ]

    CORE_EDITABLE_FIELDS = {
        "Author": "author",
        "Title": "title",
        "Subject": "subject",
        "Comments": "comments",
        "Category": "category",
        "Keywords": "keywords",
        "Language": "language",
        "Identifier": "identifier",
        "Content Status": "content_status",
        "Version": "version",
        "Last Modified By": "last_modified_by",
        "Revision": "revision",
        "Created": "created",
        "Last Modified": "modified",
    }

    EXTENDED_EDITABLE_FIELDS = {
        "Template",
        "TotalTime",
        "Application",
        "Company",
        "Manager",
        "AppVersion",
        "DocSecurity",
        "ScaleCrop",
        "LinksUpToDate",
        "SharedDoc",
        "HyperlinksChanged",
    }

    EXTENDED_CALCULATED_FIELDS = {
        "Pages",
        "Words",
        "Characters",
        "CharactersWithSpaces",
        "Lines",
        "Paragraphs",
    }

    def __init__(self, file_path):
        self.file_path = file_path
        self.metadata = {}
        self.file_info = {}
        self.core_metadata = {}
        self.extended_metadata = {}
        self.custom_metadata = {}
        self.statistics = {}

    def read_metadata(self):
        if not os.path.exists(self.file_path):
            return False

        try:
            document = Document(self.file_path)
            properties = document.core_properties

            self.file_info = self._read_file_info()

            self.statistics = {
                "Paragraph count": len(document.paragraphs),
                "Table count": len(document.tables),
                "Section count": len(document.sections),
            }

            self.core_metadata = {
                name: self._format_value(getattr(properties, attr, None))
                for name, attr in self.CORE_EDITABLE_FIELDS.items()
            }

            self.extended_metadata = self._read_app_properties()
            self.custom_metadata = self._read_custom_properties()

            self.metadata = {}
            self.metadata.update(self.file_info)
            self.metadata.update(self.core_metadata)
            self.metadata.update(self.extended_metadata)
            self.metadata.update(self.custom_metadata)

            return True

        except Exception as error:
            print(f"[!] DOCX read error: {error}")
            return False

    def _read_file_info(self):
        stat = os.stat(self.file_path)

        return {
            "FileName": os.path.basename(self.file_path),
            "FileSize": self._format_file_size(stat.st_size),
            "FileModifyDate": self._format_timestamp(stat.st_mtime),
            "FileAccessDate": self._format_timestamp(stat.st_atime),
            "FileInodeChangeDate": self._format_timestamp(stat.st_ctime),
            "FileType": "DOCX",
            "FileTypeExtension": "docx",
            "MIMEType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

    def _format_file_size(self, size_bytes):
        if size_bytes < 1024:
            return f"{size_bytes} B"

        if size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} kB"

        return f"{size_bytes / (1024 * 1024):.1f} MB"

    def _format_timestamp(self, timestamp):
        return datetime.fromtimestamp(timestamp).strftime("%Y-%m-%d %H:%M:%S")

    def _parse_datetime(self, value):
        value = str(value).strip()

        formats = [
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%dT%H:%M:%S",
            "%Y-%m-%d",
        ]

        for fmt in formats:
            try:
                parsed = datetime.strptime(value, fmt)

                if fmt == "%Y-%m-%d":
                    parsed = parsed.replace(hour=0, minute=0, second=0)

                return parsed

            except ValueError:
                pass

        try:
            return datetime.fromisoformat(value)
        except ValueError:
            raise ValueError("Invalid date format. Use YYYY-MM-DD HH:MM:SS")

    def _format_value(self, value):
        return "" if value is None else str(value)

    def _clean_xml_tag(self, tag):
        return tag.split("}")[-1]

    def _read_app_properties(self):
        result = {}

        try:
            with zipfile.ZipFile(self.file_path, "r") as docx_zip:
                if "docProps/app.xml" not in docx_zip.namelist():
                    return result

                xml_data = docx_zip.read("docProps/app.xml")
                root = ET.fromstring(xml_data)

                for element in root:
                    tag_name = self._clean_xml_tag(element.tag)
                    result[tag_name] = element.text or ""

        except Exception as error:
            result["Extended metadata error"] = str(error)

        return result

    def _read_custom_properties(self):
        result = {}

        try:
            with zipfile.ZipFile(self.file_path, "r") as docx_zip:
                if "docProps/custom.xml" not in docx_zip.namelist():
                    return result

                xml_data = docx_zip.read("docProps/custom.xml")
                root = ET.fromstring(xml_data)

                for prop in root:
                    property_name = prop.attrib.get("name", "UnknownCustomProperty")
                    property_value = ""

                    for child in prop:
                        property_value = child.text or ""

                    result[f"Custom:{property_name}"] = property_value

        except Exception as error:
            result["Custom metadata error"] = str(error)

        return result

    def get_metadata(self):
        return {
            "file_path": self.file_path,
            "file_name": os.path.basename(self.file_path),
            "statistics": self.file_info,
            "technical_metadata": self.statistics,
            "editable_metadata": self.core_metadata,
            "extended_metadata": self.extended_metadata,
            "custom_metadata": self.custom_metadata,
            "all_metadata": self.metadata,
        }

    def get_editable_tags(self):
        tags = []

        tags.extend(self.FILESYSTEM_EDITABLE_FIELDS)
        tags.extend(list(self.CORE_EDITABLE_FIELDS.keys()))

        for tag in self.extended_metadata.keys():
            if tag in self.EXTENDED_EDITABLE_FIELDS:
                tags.append(f"Extended:{tag}")

        return tags

    def get_tags(self):
        return list(self.metadata.keys())

    def edit_metadata(self, tag, new_value):
        if tag in self.FILESYSTEM_EDITABLE_FIELDS:
            return self.edit_filesystem_metadata(tag, new_value)

        if tag.startswith("Extended:"):
            extended_tag = tag.replace("Extended:", "", 1)
            return self.edit_extended_metadata(extended_tag, new_value)

        return self.edit_core_metadata(tag, new_value)

    def edit_filesystem_metadata(self, tag, new_value):
        try:
            parsed_datetime = self._parse_datetime(new_value)
            new_timestamp = parsed_datetime.timestamp()

            stat = os.stat(self.file_path)
            current_access_time = stat.st_atime
            current_modify_time = stat.st_mtime

            if tag == "FileModifyDate":
                os.utime(self.file_path, (current_access_time, new_timestamp))

            elif tag == "FileAccessDate":
                os.utime(self.file_path, (new_timestamp, current_modify_time))

            else:
                print(f"[!] Filesystem field '{tag}' is not editable.")
                return False

            self.read_metadata()
            print(f"[+] Updated filesystem field '{tag}'.")
            return True

        except Exception as error:
            print(f"[!] Filesystem metadata write error: {error}")
            return False

    def edit_core_metadata(self, tag, new_value):
        if tag not in self.CORE_EDITABLE_FIELDS:
            print(f"[!] Field '{tag}' cannot be edited as core metadata.")
            return False

        try:
            document = Document(self.file_path)
            properties = document.core_properties

            attr = self.CORE_EDITABLE_FIELDS[tag]
            converted_value = self._convert_core_value(tag, new_value)

            setattr(properties, attr, converted_value)
            document.save(self.file_path)

            self.read_metadata()
            print(f"[+] Updated core field '{tag}'.")
            return True

        except Exception as error:
            print(f"[!] DOCX core write error: {error}")
            return False

    def edit_extended_metadata(self, tag, new_value):
        if tag in self.EXTENDED_CALCULATED_FIELDS:
            print(f"[!] Field '{tag}' is calculated by Word and may be overwritten.")

        if tag not in self.extended_metadata and tag not in self.EXTENDED_EDITABLE_FIELDS:
            print(f"[!] Unknown or unsupported extended field: {tag}")
            return False

        temp_fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(temp_fd)

        try:
            with zipfile.ZipFile(self.file_path, "r") as source_zip:
                with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target_zip:
                    for item in source_zip.infolist():
                        data = source_zip.read(item.filename)

                        if item.filename == "docProps/app.xml":
                            data = self._modify_app_xml(data, tag, new_value)

                        target_zip.writestr(item, data)

            shutil.move(temp_path, self.file_path)
            self.read_metadata()

            print(f"[+] Updated extended field '{tag}'.")
            return True

        except Exception as error:
            if os.path.exists(temp_path):
                os.remove(temp_path)

            print(f"[!] DOCX extended write error: {error}")
            return False

    def _modify_app_xml(self, xml_data, tag, new_value):
        ET.register_namespace(
            "",
            "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
        )
        ET.register_namespace(
            "vt",
            "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
        )

        root = ET.fromstring(xml_data)

        for element in root:
            tag_name = self._clean_xml_tag(element.tag)

            if tag_name == tag:
                element.text = str(new_value)
                return ET.tostring(root, encoding="utf-8", xml_declaration=True)

        namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
        new_element = ET.SubElement(root, f"{{{namespace}}}{tag}")
        new_element.text = str(new_value)

        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _convert_core_value(self, tag, value):
        if tag in ["Created", "Last Modified"]:
            if isinstance(value, datetime):
                return value

            return self._parse_datetime(value)

        if tag == "Revision":
            return int(value)

        return value

    def remove_tag(self, tag):
        if tag in self.FILESYSTEM_EDITABLE_FIELDS:
            print("[!] Filesystem timestamps cannot be removed, only modified.")
            return False

        if tag.startswith("Extended:"):
            extended_tag = tag.replace("Extended:", "", 1)
            return self.edit_extended_metadata(extended_tag, "")

        if tag == "Revision":
            return self.edit_core_metadata(tag, "1")

        return self.edit_core_metadata(tag, "")

    def clear_all_editable_metadata(self):
        results = {}

        for tag in self.get_editable_tags():
            if tag in [
                "FileModifyDate",
                "FileAccessDate",
                "Created",
                "Last Modified",
                "Revision",
            ]:
                continue

            results[tag] = self.remove_tag(tag)

        return results

    def display_metadata(self):
        print("\n" + "=" * 60)
        print(f"  DOCX REPORT: {os.path.basename(self.file_path)}")
        print("=" * 60)

        print("[ File Information ]")
        for key, value in self.file_info.items():
            print(f"  > {key:25}: {value}")

        print("\n[ File Statistics ]")
        for key, value in self.statistics.items():
            print(f"  > {key:25}: {value}")

        print("\n[ Core Properties ]")
        for key, value in self.core_metadata.items():
            print(f"  > {key:25}: {value}")

        print("\n[ Extended Office Properties ]")
        for key, value in self.extended_metadata.items():
            print(f"  > {key:25}: {value}")

        print("\n[ Custom Properties ]")
        if not self.custom_metadata:
            print("  No custom metadata found.")
        else:
            for key, value in self.custom_metadata.items():
                print(f"  > {key:25}: {value}")

        print("=" * 60)